from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook

from app.parsers import ParserRegistry
from app.parsers.docx_parser import DocxParser
from app.parsers.pdf_parser import PdfParser
from app.parsers.source_parser import SourceParser
from app.parsers.xlsx_parser import XlsxParser

PARSER_SAMPLE_DIR = Path(__file__).resolve().parents[1] / "storage" / "parser_tests"


def _sample_path(filename: str) -> Path:
    PARSER_SAMPLE_DIR.mkdir(parents=True, exist_ok=True)
    return PARSER_SAMPLE_DIR / filename


def _pdf_bytes(text: str) -> bytes:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET\n".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"endstream",
    ]

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode("ascii"))
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")

    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(offsets)}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        (
            f"trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(pdf)


def test_xlsx_parser_extracts_all_sheets_with_coordinates() -> None:
    path = _sample_path("parser_sample.xlsx")
    workbook = Workbook()
    first_sheet = workbook.active
    first_sheet.title = "First"
    first_sheet["A1"] = "ID"
    first_sheet["B2"] = "Value"
    second_sheet = workbook.create_sheet("Second")
    second_sheet["A1"] = "Other"
    workbook.save(path)

    parsed = XlsxParser().parse(path)

    assert "# Sheet: First" in parsed.content
    assert "First!A1: ID" in parsed.content
    assert "First!B2: Value" in parsed.content
    assert "# Sheet: Second" in parsed.content
    assert "Second!A1: Other" in parsed.content


def test_docx_parser_extracts_paragraphs_and_tables() -> None:
    path = _sample_path("parser_sample.docx")
    document = Document()
    document.add_paragraph("Paragraph text")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cell A"
    table.cell(0, 1).text = "Cell B"
    document.save(path)

    parsed = DocxParser().parse(path)

    assert "Paragraph 1: Paragraph text" in parsed.content
    assert "Table 1 R1C1: Cell A" in parsed.content
    assert "Table 1 R1C2: Cell B" in parsed.content


def test_pdf_parser_extracts_text() -> None:
    path = _sample_path("parser_sample.pdf")
    path.write_bytes(_pdf_bytes("PDF parser text"))

    parsed = PdfParser().parse(path)

    assert "PDF parser text" in parsed.content


def test_source_parser_adds_line_numbers() -> None:
    path = _sample_path("parser_sample.py")
    path.write_text("first\nsecond\n", encoding="utf-8")

    parsed = SourceParser().parse(path)

    assert parsed.content == "1: first\n2: second"


def test_parser_registry_dispatches_by_extension_and_file_type() -> None:
    path = _sample_path("registry_sample.py")
    path.write_text("print('ok')\n", encoding="utf-8")
    registry = ParserRegistry()

    by_extension = registry.parse(path)
    by_file_type = registry.parse(path, "source")

    assert by_extension.content == "1: print('ok')"
    assert by_file_type.content == "1: print('ok')"
