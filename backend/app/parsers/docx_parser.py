from __future__ import annotations

from pathlib import Path

from docx import Document

from app.parsers.base import ParsedArtifact


class DocxParser:
    def parse(self, path: Path) -> ParsedArtifact:
        document = Document(path)
        lines: list[str] = []
        for index, paragraph in enumerate(document.paragraphs, start=1):
            if paragraph.text:
                lines.append(f"Paragraph {index}: {paragraph.text}")

        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                for cell_index, cell in enumerate(row.cells, start=1):
                    lines.append(f"Table {table_index} R{row_index}C{cell_index}: {cell.text}")
        return ParsedArtifact(content="\n".join(lines))
